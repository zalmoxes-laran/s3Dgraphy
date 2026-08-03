"""BK1 — a baked narrative rendered to .docx.

**The DocX is for the normal user.** LaTeX (L1) is for the apparatus and Jupyter
for the analysis; a Word file is what an archaeologist sends to a colleague, a
superintendency or an editor. So this renderer optimises for something none of the
others do: being *opened and read* by someone who will never see the graph.

**It renders a bake, it does not traverse a graph.** Everything about what the
narrative says was decided in :mod:`s3dgraphy.narrative.bake`; this module only
places it. That is why there is no `view_type` switch here — the bake already
turned the vocabulary into `kind`s a page can hold, and a second interpretation
would be a second chance to disagree.

**python-docx is optional and lazy**, like rdflib and pyproj: importing
s3dgraphy, loading a graph or baking a narrative must not require it. Absent, the
caller gets :class:`~s3dgraphy.api.MissingDependency` — a "this build cannot",
which a transport maps to 501 and a client can degrade from honestly.
"""

from __future__ import annotations

import io
from typing import Any, List

from ..narrative.bake import BakedNarrative


def _require_docx():
    """Import python-docx or say precisely what is missing.

    Kept in one place so every entry point fails the same way, with a message
    that names the package AND how to get it — an ImportError that only says
    'docx' sends the reader to the wrong project (there is an unrelated `docx` on
    PyPI; the one we need is `python-docx`).
    """
    try:
        import docx  # noqa: F401
        return docx
    except ImportError as exc:
        from ..api import MissingDependency
        raise MissingDependency(
            "DocX export needs python-docx — install the extra: "
            "pip install 's3dgraphy[docx]' (the PyPI package is "
            "'python-docx', not 'docx')"
        ) from exc


def _add_citation_paragraph(document: Any, entry: dict) -> None:
    """One bibliography line: author, year, title, publisher, url — whichever
    exist. Fields the source does not answer are skipped, never filled in: a
    citable invented year is the one error a bibliography must not make."""
    bits: List[str] = []
    if entry.get("author"):
        bits.append(str(entry["author"]))
    if entry.get("year"):
        bits.append(f"({entry['year']})")
    bits.append(str(entry.get("title", "")))
    if entry.get("publisher"):
        bits.append(str(entry["publisher"]))
    if entry.get("url"):
        bits.append(str(entry["url"]))
    document.add_paragraph(", ".join(b for b in bits if b), style=None)


def render_docx(baked: BakedNarrative) -> bytes:
    """Render a :class:`BakedNarrative` to .docx bytes.

    Returns bytes rather than writing a file: the caller may be a CLI, an HTTP
    response or a notebook, and only one of those wants a path.
    """
    docx = _require_docx()
    from docx.shared import Inches

    document = docx.Document()

    document.add_heading(baked.title, level=0)
    if baked.description:
        document.add_paragraph(baked.description)

    # ── byline (N8): responsibility and assistance are different claims ───────
    # A person can be asked about a sentence; a model cannot. So people get the
    # byline and models get a stated assistance line — never the same line, which
    # would attribute accountability to something that has none.
    if baked.responsible:
        paragraph = document.add_paragraph()
        run = paragraph.add_run("A cura di " + ", ".join(baked.responsible))
        run.italic = True
    if baked.assisting:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(
            "Con l'assistenza di " + ", ".join(baked.assisting)
            + ". Il testo generato è stato rivisto e avallato da una persona, "
              "tranne dove indicato.")
        run.italic = True
        run.font.size = docx.shared.Pt(9)

    if baked.pending_validation:
        count = baked.pending_validation
        paragraph = document.add_paragraph()
        run = paragraph.add_run(
            f"Nota. {count} "
            + ("paragrafo di questo testo è" if count == 1
               else "paragrafi di questo testo sono")
            + " una bozza generata automaticamente e non ancora avallata; "
              "sono segnalati nel testo.")
        run.bold = True

    for chapter in baked.chapters:
        document.add_heading(chapter.title, level=1)
        for block in chapter.blocks:
            if block.kind == "prose":
                paragraph = document.add_paragraph()
                if block.unendorsed:
                    # In the text, not in a comment or a style: a Word file gets
                    # copy-pasted, printed and re-saved, and every one of those
                    # loses anything that is not a character.
                    flag = paragraph.add_run("[bozza generata, non avallata] ")
                    flag.italic = True
                    flag.bold = True
                paragraph.add_run(block.text)
                continue

            if block.kind == "image":
                if block.image is not None and block.image.resolved:
                    document.add_picture(io.BytesIO(block.image.data),
                                         width=Inches(5.5))
                    caption = document.add_paragraph(block.text)
                    caption.runs[0].italic = True if caption.runs else None
                else:
                    note = (block.image.note if block.image is not None
                            else "no image")
                    paragraph = document.add_paragraph()
                    run = paragraph.add_run(
                        f"[immagine non incorporata: {block.text} — {note}]")
                    run.italic = True
                continue

            if block.kind == "citation":
                paragraph = document.add_paragraph()
                paragraph.add_run(block.text)
                entry = block.meta.get("citation") or {}
                if entry.get("author") or entry.get("year"):
                    detail = " (" + ", ".join(
                        str(entry[k]) for k in ("author", "year")
                        if entry.get(k)) + ")"
                    paragraph.add_run(detail)
                # A document that IS a picture shows it under the citation.
                if block.image is not None and block.image.resolved:
                    document.add_picture(io.BytesIO(block.image.data),
                                         width=Inches(5.0))
                continue

            if block.kind == "map":
                paragraph = document.add_paragraph()
                paragraph.add_run(block.text)
                if block.link:
                    paragraph.add_run(f" — {block.link}")
                continue

            if block.kind == "placeholder":
                paragraph = document.add_paragraph()
                run = paragraph.add_run(f"[{block.text}]")
                run.italic = True
                if block.link:
                    paragraph.add_run(f" {block.link}")
                continue

            if block.kind == "unresolved":
                paragraph = document.add_paragraph()
                run = paragraph.add_run(block.text)
                run.bold = True
                continue

            # `unit` and anything the bake adds later: plain text. A renderer
            # that dropped an unknown kind would remove evidence from a published
            # document, where nobody can notice it is gone.
            document.add_paragraph(block.text)

    if baked.citations:
        document.add_heading("Fonti", level=1)
        for entry in baked.citations:
            _add_citation_paragraph(document, entry)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()

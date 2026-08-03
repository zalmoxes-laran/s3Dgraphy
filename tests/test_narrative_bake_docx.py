"""BK1 — the bake, and the .docx rendered from it.

A bake is a **snapshot**: the moment a living narrative becomes a thing you can
send someone. What these tests defend is not "the output has paragraphs" but the
four ways a snapshot betrays the text it froze:

* **a hole gets hidden.** An embed that will not resolve must become a block that
  SAYS so and appear in ``baked.unresolved``. Raising would throw away a whole
  publication over one missing file; dropping the block removes evidence from a
  document where nobody can notice it is gone.
* **certainty is lost.** A stratigraphic unit carries a hypothesis. A static
  projection that prints the label and drops the certainty turns a hypothesis into
  a statement — the most consequential thing this projection can get wrong.
* **an unendorsed machine draft reads like endorsed prose.** A .docx is printed,
  copy-pasted and re-saved, and each of those loses anything that is not a
  character. So the flag has to be IN the text.
* **the byline credits a model as an author.** People can be asked about a claim;
  a model cannot. The two lists must stay apart.

Plus the property that makes the arrangement worth having: the bake is
**independent of every renderer**, so it must work with python-docx absent.
"""

from __future__ import annotations

import io
import os
import sys
import zipfile

import pytest

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "src"))

from s3dgraphy import api  # noqa: E402
from s3dgraphy.graph import Graph  # noqa: E402
from s3dgraphy.nodes.document_node import DocumentNode  # noqa: E402
from s3dgraphy.nodes.narrative_node import NarrativeNode  # noqa: E402
from s3dgraphy.nodes.stratigraphic_node import StratigraphicUnit  # noqa: E402

FIXTURE = "tests/fixtures/PortaMarina-lite.em.json"
FIXTURE_DIR = "tests/fixtures"

docx_missing = pytest.mark.skipif(
    __import__("importlib").util.find_spec("docx") is None,
    reason="python-docx not installed (the [docx] extra)")

#: The smallest real PNG: 1x1, opaque. A byte literal rather than a committed
#: binary, so the fixture is legible in a diff and cannot go missing.
_PNG_1x1 = bytes.fromhex(
    "89504e470d0a1a0a0000000d49484452000000010000000108060000001f15c4"
    "890000000a49444154789c6300010000050001od0a2db40000000049454e44ae"
    "426082".replace("od", "7d"))


def _fixture_graph():
    graph, _warnings = api.load_emjson_file(FIXTURE)
    return graph


def _image_graph(tmp_path, *, write_file=True, locator=None):
    """A minimal narrative with a picture embed and a source, per BK1's ask.

    The picture is an **`rm`** and not an "image": the EM vocabulary has no
    `image` view_type (a picture arrives as a `document` or as a representation
    model), and `Block` rejects an unknown one at construction. This fixture is
    therefore also the proof that the bake speaks the datamodel's language.

    Built here rather than added to PortaMarina-lite: that fixture is compared
    byte-for-byte by two other tests, so growing it to serve this one would make
    an unrelated suite fail.
    """
    png = tmp_path / "plan.png"
    if write_file:
        png.write_bytes(_PNG_1x1)

    graph = Graph(graph_id="baked")
    plan = DocumentNode(node_id="D.plan", name="Pianta di scavo",
                        description="Rilievo 1:50",
                        url=locator if locator is not None else "plan.png")
    plan.data.update({"author": "Roe, Jane", "year": "2024"})
    graph.add_node(plan)
    source = DocumentNode(node_id="D.report", name="Relazione di scavo")
    source.data.update({"author": "Doe, John", "year": "2019"})
    graph.add_node(source)
    unit = StratigraphicUnit(node_id="US.1", name="US1",
                             description="Muro in opera incerta")
    unit.data = {"certainty": "ipotetica"}
    graph.add_node(unit)

    narrative = NarrativeNode(node_id="N.1", name="Storia")
    graph.add_node(narrative)
    chapter = narrative.add_chapter("Capitolo")
    chapter.add_prose("Prosa introduttiva.")
    chapter.add_embed("D.plan", "rm", caption="La pianta")
    chapter.add_embed("D.report", "source")
    chapter.add_embed("US.1", "us")
    return graph


# ── the bake: what it resolves ────────────────────────────────────────────────

def test_a_source_bakes_into_a_citation_with_a_stable_key():
    """The key is the same one L1 puts in the `.bib`, so a DocX bibliography and
    a LaTeX one cannot disagree about what was cited."""
    baked = api.bake_narrative(_fixture_graph(), "NARR.portamarina",
                               base_dir=FIXTURE_DIR)
    keys = [c["key"] for c in baked.citations]
    assert api.bib_key("D.1") in keys
    assert len(keys) == len(set(keys)), "a source cited twice must yield one entry"


def test_the_bake_reads_image_bytes(tmp_path):
    graph = _image_graph(tmp_path)
    baked = api.bake_narrative(graph, "N.1", base_dir=str(tmp_path))
    images = [b for b in baked.chapters[0].blocks if b.kind == "image"]
    assert len(images) == 1
    assert images[0].image is not None and images[0].image.resolved
    assert images[0].image.data == _PNG_1x1
    assert images[0].text == "La pianta"


def test_a_missing_image_becomes_a_stated_hole_not_an_exception(tmp_path):
    """The whole reason the bake never raises: one absent file must not cost a
    publication. But the hole is recorded twice — inline, so a reader of the
    output sees it, and in ``unresolved``, so a script can refuse to publish."""
    graph = _image_graph(tmp_path, write_file=False)
    baked = api.bake_narrative(graph, "N.1", base_dir=str(tmp_path))
    image = next(b for b in baked.chapters[0].blocks if b.kind == "image")
    assert not (image.image and image.image.resolved)
    assert "not found" in (image.image.note if image.image else "")
    assert "D.plan" in baked.unresolved


def test_an_unresolvable_reference_is_never_silently_dropped(tmp_path):
    """A ref to a node the graph no longer holds. Dropping it would delete
    evidence from a text where its absence is undetectable."""
    graph = _image_graph(tmp_path)
    narrative = graph.find_node_by_id("N.1")
    narrative.chapters[0].add_embed("D.gone", "source")
    baked = api.bake_narrative(graph, "N.1", base_dir=str(tmp_path))
    holes = [b for b in baked.chapters[0].blocks if b.kind == "unresolved"]
    assert len(holes) == 1
    assert "D.gone" in holes[0].text
    assert "D.gone" in baked.unresolved


def test_the_bake_does_not_fetch_remote_images(tmp_path):
    """Offline by construction. Fetching would make publishing depend on the
    network at bake time, and would let a remote change alter a snapshot that is
    supposed to be frozen."""
    graph = _image_graph(tmp_path, write_file=False,
                         locator="https://example.org/plan.png")
    baked = api.bake_narrative(graph, "N.1", base_dir=str(tmp_path))
    image = next(b for b in baked.chapters[0].blocks if b.kind == "image")
    assert not (image.image and image.image.resolved)
    assert "not fetched" in image.image.note


def test_certainty_survives_the_bake(tmp_path):
    """The one that matters most. A unit's certainty printed away turns a
    hypothesis into a statement, and the reader of the .docx has no way back to
    the graph to find out."""
    graph = _image_graph(tmp_path)
    baked = api.bake_narrative(graph, "N.1", base_dir=str(tmp_path))
    unit = next(b for b in baked.chapters[0].blocks if b.kind == "unit")
    assert "ipotetica" in unit.text


def test_a_map_bakes_to_coordinates_and_a_link():
    """A tile render is a follow-up; the POSITION can be frozen honestly, so it
    is. The link is what a reader can act on without the graph."""
    baked = api.bake_narrative(_fixture_graph(), "NARR.portamarina",
                               base_dir=FIXTURE_DIR)
    maps = [b for chapter in baked.chapters for b in chapter.blocks
            if b.kind == "map"]
    assert maps, "the fixture has a map embed"
    block = maps[0]
    assert "lat" in block.meta and "lon" in block.meta
    assert "openstreetmap.org" in block.link
    # The coordinates are IN the text too: a reader of a printed page has no link.
    assert "40.7497" in block.text


def test_deferred_renders_become_labelled_placeholders(tmp_path):
    """`scene3d` and `matrix` need renderers this build has not got. A labelled
    placeholder says which and why; a fabricated caption would claim a picture
    exists."""
    graph = _image_graph(tmp_path)
    narrative = graph.find_node_by_id("N.1")
    narrative.chapters[0].add_embed("US.1", "scene3d")
    narrative.chapters[0].add_embed("US.1", "matrix")
    baked = api.bake_narrative(graph, "N.1", base_dir=str(tmp_path))
    placeholders = [b for b in baked.chapters[0].blocks
                    if b.kind == "placeholder"]
    assert {b.view_type for b in placeholders} == {"scene3d", "matrix"}
    for block in placeholders:
        assert block.view_type in block.text
        assert "not produced by this build" in block.meta["reason"]


def test_a_vocabulary_type_with_no_static_form_is_baked_as_text_not_dropped(
        tmp_path):
    """`timeline` and `table` are in the EM vocabulary but have no static form of
    their own yet. They must still arrive in the output — as text, marked as
    tolerated rather than handled. A block that silently vanished from a published
    document is unrecoverable, and the note is what keeps the gap findable.

    (A view_type OUTSIDE the vocabulary cannot be tested: `Block.__post_init__`
    rejects it at construction, which is why the bake has no branch for one.)
    """
    graph = _image_graph(tmp_path)
    chapter = graph.find_node_by_id("N.1").chapters[0]
    chapter.add_embed("US.1", "timeline")
    chapter.add_embed("US.1", "table")
    baked = api.bake_narrative(graph, "N.1", base_dir=str(tmp_path))
    tolerated = [b for b in baked.chapters[0].blocks
                 if b.view_type in ("timeline", "table")]
    assert len(tolerated) == 2
    for block in tolerated:
        assert block.kind == "unit"
        assert block.view_type in block.meta.get("note", "")


def test_the_vocabulary_is_read_from_the_datamodel_not_restated():
    """ADR-001: no hand-written list of EM types. The bake's own categories must
    be subsets of the datamodel's vocabulary, so a type renamed there cannot leave
    a stale spelling behind here."""
    from s3dgraphy.narrative import bake as bake_module
    from s3dgraphy.nodes.narrative_node import NARRATIVE_VIEW_TYPES

    assert set(bake_module.DEFERRED_RENDER_VIEW_TYPES) <= set(
        NARRATIVE_VIEW_TYPES)
    assert set(bake_module.PICTURE_VIEW_TYPES) <= set(NARRATIVE_VIEW_TYPES)
    assert set(bake_module.CITED_VIEW_TYPES) <= set(NARRATIVE_VIEW_TYPES)
    # And "image" is NOT one of them — the type BK1's brief named does not exist.
    assert "image" not in NARRATIVE_VIEW_TYPES


def test_the_byline_keeps_people_apart_from_models():
    """N8. A model cannot be asked about a sentence, so it is never listed where
    a person is."""
    baked = api.bake_narrative(_fixture_graph(), "NARR.portamarina",
                               base_dir=FIXTURE_DIR)
    assert any("Demetrescu" in name for name in baked.responsible)
    assert any("Claude" in name for name in baked.assisting)
    assert not any("Claude" in name for name in baked.responsible)


def test_unendorsed_prose_is_marked_in_the_bake():
    baked = api.bake_narrative(_fixture_graph(), "NARR.portamarina",
                               base_dir=FIXTURE_DIR)
    prose = [b for chapter in baked.chapters for b in chapter.blocks
             if b.kind == "prose"]
    assert any(b.unendorsed for b in prose)
    assert baked.pending_validation == sum(1 for b in prose if b.unendorsed)


def test_the_exporter_is_the_only_source_of_the_unendorsed_marker():
    """The flag comes from the STRUCTURE (`ai_generated` and no `validated_by`),
    never from the prose.

    The fixture used to carry `[bozza non avallata]` typed into the block's text
    *as well*, so both L1 and the DocX printed the warning twice — belt and braces
    that read as a bug to anybody holding the page. Removing it leaves one source
    of truth, and this test guards both directions: a hand-typed marker creeping
    back into a fixture, and an exporter quietly dropping its own.

    Deliberately checks the LaTeX projection too, although this file is BK1's:
    the fixture is shared, and a guard that only covered one renderer would let
    the other drift.
    """
    graph = _fixture_graph()
    narrative = graph.find_node_by_id("NARR.portamarina")
    prose = [b for chapter in narrative.chapters
             for b in chapter.blocks
             if getattr(b, "block_type", "") == "prose"]
    assert prose, "the fixture has prose"
    for block in prose:
        assert "avallat" not in (block.text or "").lower(), (
            "a block's TEXT must not carry the unendorsed marker — the exporter "
            "adds it from ai_generated/validated_by")

    tex = api.export_narrative_latex(graph, "NARR.portamarina")["tex"]
    assert tex.count("bozza generata, non avallata") == 1


def test_baking_an_unknown_narrative_raises():
    """Baking nothing under a name the caller believes in is worse than failing."""
    with pytest.raises(KeyError):
        api.bake_narrative(_fixture_graph(), "N.does-not-exist")


# ── the degradation that makes the split worth having ─────────────────────────

def test_the_bake_does_not_need_the_docx_renderer(monkeypatch):
    """The bake is independent of every renderer — that is the design claim, so
    it is asserted with python-docx made unimportable."""
    monkeypatch.setitem(sys.modules, "docx", None)
    baked = api.bake_narrative(_fixture_graph(), "NARR.portamarina",
                               base_dir=FIXTURE_DIR)
    assert baked.chapters and baked.citations


def test_docx_without_python_docx_is_a_clean_MissingDependency(monkeypatch):
    """Not an ImportError from somewhere deep: a transport maps MissingDependency
    to 501 ('this build cannot'), which a client can degrade from. And the
    message names `python-docx`, because there is an unrelated `docx` on PyPI and
    installing it would not help."""
    import importlib

    real_find_spec = importlib.util.find_spec

    def hide_docx(name, *args, **kwargs):
        if name == "docx":
            return None
        return real_find_spec(name, *args, **kwargs)

    monkeypatch.setattr(importlib.util, "find_spec", hide_docx)
    monkeypatch.setitem(sys.modules, "docx", None)
    with pytest.raises(api.MissingDependency) as exc:
        api.export_narrative_docx(_fixture_graph(), "NARR.portamarina")
    assert "python-docx" in str(exc.value)


# ── the .docx itself ─────────────────────────────────────────────────────────

@docx_missing
def test_the_docx_is_a_document_word_can_open():
    """Not "the bytes are non-empty": a .docx is a zip of OOXML parts, and a
    renderer can emit something plausible that no reader opens. So it is reopened
    and read back."""
    import docx

    blob = api.export_narrative_docx(_fixture_graph(), "NARR.portamarina",
                                     base_dir=FIXTURE_DIR)
    assert "word/document.xml" in zipfile.ZipFile(io.BytesIO(blob)).namelist()
    document = docx.Document(io.BytesIO(blob))
    texts = [p.text for p in document.paragraphs]
    assert any("Porta Marina" in t for t in texts)
    # Every chapter title, as a heading rather than as bold prose: a heading is
    # what gives the reader a navigable document.
    headings = [p.text for p in document.paragraphs
                if p.style.name.startswith("Heading")]
    assert "Presentazione" in headings
    assert "Età imperiale" in headings


@docx_missing
def test_the_docx_embeds_the_image(tmp_path):
    """The picture has to be IN the file — a .docx that references a path on the
    author's disk arrives blank at the recipient, which is the whole failure the
    bake's byte-reading exists to prevent."""
    graph = _image_graph(tmp_path)
    blob = api.export_narrative_docx(graph, "N.1", base_dir=str(tmp_path))
    names = zipfile.ZipFile(io.BytesIO(blob)).namelist()
    media = [n for n in names if n.startswith("word/media/")]
    assert media, f"no embedded media in the .docx: {names}"


@docx_missing
def test_the_docx_says_in_words_that_a_draft_is_unendorsed():
    """On paper there is no badge. If this text is not there, an unreviewed
    machine draft is typographically identical to endorsed prose."""
    import docx

    blob = api.export_narrative_docx(_fixture_graph(), "NARR.portamarina",
                                     base_dir=FIXTURE_DIR)
    body = "\n".join(p.text for p in docx.Document(io.BytesIO(blob)).paragraphs)
    assert "bozza generata, non avallata" in body
    # And the count is stated up front, so the reader knows before reading.
    assert "non ancora avallata" in body


@docx_missing
def test_the_docx_carries_its_sources_and_the_prompt():
    """A generated passage's prompt is a source like any other: 'how did the
    machine come to write this' must be answerable from the file alone."""
    import docx

    blob = api.export_narrative_docx(_fixture_graph(), "NARR.portamarina",
                                     base_dir=FIXTURE_DIR)
    body = "\n".join(p.text for p in docx.Document(io.BytesIO(blob)).paragraphs)
    assert "Fonti" in body
    assert "Rilievo Maiuri" in body
    assert "Prompt" in body


@docx_missing
def test_the_cli_writes_a_docx(tmp_path):
    """The CLI is part of the surface, so it is exercised rather than assumed."""
    out = tmp_path / "out.docx"
    code = api.main(["export-narrative-docx", FIXTURE, "NARR.portamarina",
                     "-o", str(out)])
    assert code == 0
    assert out.exists() and out.stat().st_size > 0
    assert "word/document.xml" in zipfile.ZipFile(out).namelist()


@docx_missing
def test_the_cli_refuses_to_overwrite_without_force(tmp_path):
    out = tmp_path / "out.docx"
    out.write_bytes(b"existing")
    assert api.main(["export-narrative-docx", FIXTURE, "-o", str(out)]) == 1
    assert out.read_bytes() == b"existing"
    assert api.main(["export-narrative-docx", FIXTURE, "-o", str(out),
                     "--force"]) == 0
    assert out.read_bytes() != b"existing"


def test_the_cli_reports_a_missing_renderer_with_its_own_exit_code(
        tmp_path, monkeypatch):
    """Exit 3 = "this build cannot", distinct from exit 1 = "your input is wrong",
    so a script can tell them apart without parsing a message.

    This also pins a trap that only appears when the CLI is run as
    ``python -m s3dgraphy.api``: this module is then ``__main__``, and a submodule
    importing ``from ..api import MissingDependency`` gets a SECOND copy of the
    class. Catching the local name would not match what the exporter raised — the
    user would get a traceback instead of exit 3, and no file, with no explanation.
    """
    import importlib

    real_find_spec = importlib.util.find_spec

    def hide_docx(name, *args, **kwargs):
        if name == "docx":
            return None
        return real_find_spec(name, *args, **kwargs)

    monkeypatch.setattr(importlib.util, "find_spec", hide_docx)
    monkeypatch.setitem(sys.modules, "docx", None)
    out = tmp_path / "out.docx"
    assert api.main(["export-narrative-docx", FIXTURE, "-o", str(out)]) == 3
    # And nothing was written: a zero-byte .docx would be worse than no file.
    assert not out.exists()


@docx_missing
def test_a_missing_image_is_stated_in_the_docx(tmp_path):
    """The hole travels into the rendered file, not only into the bake's list."""
    import docx

    graph = _image_graph(tmp_path, write_file=False)
    blob = api.export_narrative_docx(graph, "N.1", base_dir=str(tmp_path))
    body = "\n".join(p.text for p in docx.Document(io.BytesIO(blob)).paragraphs)
    assert "immagine non incorporata" in body
    assert "La pianta" in body
